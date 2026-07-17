#!/usr/bin/env python3
"""Build deterministic Word team-reference copies of policy-safe recut scripts."""

from __future__ import annotations

import json
import hashlib
import re
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PACKAGE_DIR = ROOT / "docs/course-production/held-video-recuts"
OUTPUT_DIR = PACKAGE_DIR / "generated"
SOURCE_KEYS = (
    "video-slot-01-welcome",
    "video-slot-01-mindset",
    "video-slot-10-objection-scripts",
    "video-slot-15-closing",
    "video-slot-17-compensation",
    "video-slot-18-operator",
    "video-slot-19-career",
)

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
DOCUMENT_TIMESTAMP = datetime(2026, 7, 17, 12, 0, 0, tzinfo=timezone.utc)
INVENTORY_PATH = OUTPUT_DIR / "team-reference-docx.json"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_run_font(run, *, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_style(style, *, size, color, before, after, line_spacing=1.25):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    field_run._r.extend((begin, instruction, separate))
    paragraph.add_run("1")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def human_title(source_key: str, package: dict) -> str:
    first_title = package["scenes"][0]["title"].strip()
    if source_key == "video-slot-01-welcome":
        return "Welcome to BMH Group"
    return first_title


def normalize_docx_archive(output_path: Path) -> None:
    """Normalize ZIP ordering and timestamps so identical inputs produce identical bytes."""
    with ZipFile(output_path, "r") as source_archive:
        entries = [
            (entry.filename, source_archive.read(entry.filename), entry.external_attr)
            for entry in source_archive.infolist()
        ]
    temporary_path = output_path.with_suffix(".docx.tmp")
    with ZipFile(temporary_path, "w", compression=ZIP_DEFLATED, compresslevel=9) as target_archive:
        for filename, contents, external_attr in sorted(entries):
            info = ZipInfo(filename, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = external_attr
            target_archive.writestr(info, contents)
    temporary_path.replace(output_path)


def build_document(package_path: Path) -> Path:
    package = json.loads(package_path.read_text(encoding="utf-8"))
    source = package["source"]
    source_key = source["source_key"]
    constraints = package["production_constraints"]
    if constraints["provider_call_allowed"] or constraints["render_allowed"]:
        raise ValueError(f"{source_key} unexpectedly authorizes provider work")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    styles = doc.styles
    set_style(styles["Normal"], size=11, color=RGBColor(0, 0, 0), before=0, after=6)
    set_style(styles["Heading 1"], size=16, color=BLUE, before=18, after=10)
    set_style(styles["Heading 2"], size=13, color=BLUE, before=14, after=7)
    set_style(styles["Heading 3"], size=12, color=DARK_BLUE, before=10, after=5)

    header = section.header.paragraphs[0]
    header.text = "BMH Institute | Replacement video script"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    add_page_field(footer)
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    set_run_font(kicker.add_run("BMH INSTITUTE VIDEO PRODUCTION"), size=9, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run(human_title(source_key, package)), size=26, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("Policy-safe replacement script | Team reference"), size=13, color=MUTED)

    metadata = (
        ("Source key", source_key),
        ("Held-source SHA-256", source["held_sha256"]),
        ("Status", "Script ready; provider call, render, captions, approval, and publication remain gated"),
        ("Authorization", "This document does not authorize a HeyGen provider call or render"),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        set_run_font(paragraph.add_run(f"{label}: "), bold=True)
        set_run_font(paragraph.add_run(value))

    doc.add_heading("Lesson contract", level=1)
    objective = doc.add_paragraph()
    set_run_font(objective.add_run("Objective: "), bold=True)
    set_run_font(objective.add_run(package["lesson_contract"]["objective"]))

    doc.add_heading("Spoken script", level=1)
    for index, scene in enumerate(package["scenes"], start=1):
        scene_title = re.sub(r"\s+", " ", scene["title"].strip())
        scene_heading = doc.add_heading(f"Scene {index}: {scene_title}", level=2)
        scene_heading.paragraph_format.keep_with_next = True
        spoken = doc.add_paragraph(scene["spoken_text"].strip())
        spoken.paragraph_format.keep_together = True
        spoken.paragraph_format.keep_with_next = True

        visual = scene["visual_plan"]
        visual_paragraph = doc.add_paragraph()
        visual_paragraph.paragraph_format.space_after = Pt(3)
        visual_paragraph.paragraph_format.keep_together = True
        visual_paragraph.paragraph_format.keep_with_next = True
        set_run_font(visual_paragraph.add_run("Visual plan: "), bold=True, color=DARK_BLUE)
        set_run_font(visual_paragraph.add_run(visual["shot"]))
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(8)
        note.paragraph_format.keep_together = True
        set_run_font(note.add_run("Editor note: "), bold=True, color=DARK_BLUE)
        set_run_font(note.add_run(visual["editor_note"]), italic=True, color=MUTED)

    doc.core_properties.title = f"{human_title(source_key, package)} - policy-safe replacement script"
    doc.core_properties.subject = "BMH Institute replacement video team reference"
    doc.core_properties.author = "BMH Institute content QA"
    doc.core_properties.keywords = "BMH Institute, video script, policy-safe replacement"
    doc.core_properties.created = DOCUMENT_TIMESTAMP
    doc.core_properties.modified = DOCUMENT_TIMESTAMP

    output_path = OUTPUT_DIR / f"{source_key}-script.docx"
    doc.save(output_path)
    normalize_docx_archive(output_path)
    return output_path


def main() -> None:
    inventory = {
        "schema_version": 1,
        "purpose": "Deterministic team-reference copies; not provider authorization",
        "provider_call_allowed": False,
        "documents": [],
    }
    for source_key in SOURCE_KEYS:
        package_path = PACKAGE_DIR / f"{source_key}.json"
        output = build_document(package_path)
        script_path = OUTPUT_DIR / f"{source_key}-script.txt"
        inventory["documents"].append(
            {
                "source_key": source_key,
                "path": output.relative_to(ROOT).as_posix(),
                "sha256": sha256(output),
                "size_bytes": output.stat().st_size,
                "package_path": package_path.relative_to(ROOT).as_posix(),
                "package_sha256": sha256(package_path),
                "script_path": script_path.relative_to(ROOT).as_posix(),
                "script_sha256": sha256(script_path),
            }
        )
    INVENTORY_PATH.write_text(
        json.dumps(inventory, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    for output in (ROOT / item["path"] for item in inventory["documents"]):
        print(output.relative_to(ROOT))
    print(INVENTORY_PATH.relative_to(ROOT))


if __name__ == "__main__":
    main()
